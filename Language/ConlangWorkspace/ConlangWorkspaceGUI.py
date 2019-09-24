from PyQt5.QtCore import (QDateTime, Qt, QTimer)
from PyQt5.QtWidgets import (QAction, QApplication, QCheckBox, QComboBox, QDateTimeEdit,
        QDial, QDialog, QFileDialog, QGridLayout, QGroupBox, QHBoxLayout, 
        QInputDialog, QLabel, QLineEdit,
        QListWidget, QListWidgetItem, QMainWindow, QMenuBar,
        QProgressBar, QPushButton, QRadioButton, QScrollBar, QSizePolicy,
        QSlider, QSpinBox, QStyleFactory, QTableWidget, QTabWidget, QTextEdit,
        QVBoxLayout, QWidget
    )
from PyQt5.QtGui import (QIntValidator)

import docx
from docx.shared import Pt

from LanguageEvolutionTools import get_random_rules

from CommandProcessor import CommandProcessor
from InflectionForm import InflectionForm
from Language import Language
from Lexeme import Lexeme
from Lexicon import Lexicon
from Morpheme import Morpheme
from OrthographyConverter import OrthographyConverter
from Rule import Rule
from Word import Word

import string
import sys


class ConlangWorkspaceGUI(QMainWindow):
    def __init__(self, language, parent=None):
        super(ConlangWorkspaceGUI, self).__init__(parent)
        self.originalPalette = QApplication.palette()

        self.language = language
        self.command_processor = CommandProcessor(
            gui=self,
            orthography_converter=OrthographyConverter()
        )

        self.setup_menu()

        mainLayout = QGridLayout()
        self.createTabWidget()
        mainLayout.addWidget(self.tabWidget)
        self.setCentralWidget(QWidget(self))
        self.centralWidget().setLayout(mainLayout)
        self.setWindowTitle("Conlang Workspace")

        self.setup_file_saving()
        self.setup_file_opening()

    def setup_menu(self):
        self.main_menu = self.menuBar()
        self.file_menu = self.main_menu.addMenu('File')
        #editMenu = mainMenu.addMenu('Edit')
        #viewMenu = mainMenu.addMenu('View')
        #searchMenu = mainMenu.addMenu('Search')
        #toolsMenu = mainMenu.addMenu('Tools')
        #helpMenu = mainMenu.addMenu('Help')

    def setup_file_opening(self):
        openFile = QAction("&Open File", self)
        openFile.setShortcut("Ctrl+O")
        openFile.setStatusTip('Open File')
        openFile.triggered.connect(self.open_file)
        self.file_menu.addAction(openFile)

    def open_file(self):
        fp = QFileDialog.getOpenFileName(self, 'Open File')
        if type(fp) is tuple:
            fp = fp[0]
        self.open_file_known_filepath(fp)

    def open_file_known_filepath(self, fp):
        with open(fp) as f:
            lines = f.readlines()
        lines = [x.strip() for x in lines]
        for command in lines:
            self.command_processor.process_command(command)

    def setup_file_saving(self):
        saveFile = QAction("&Save File", self)
        saveFile.setShortcut("Ctrl+S")
        saveFile.setStatusTip('Save File')
        saveFile.triggered.connect(self.save_file)
        self.file_menu.addAction(saveFile)

    def save_file(self):
        fp = QFileDialog.getSaveFileName(self, 'Save File')
        if type(fp) is tuple:
            fp = fp[0]
        contents = ""
        for command in self.command_processor.command_history:
            contents += command + "\n"
        with open(fp, 'w') as f:
            f.write(contents)

    def createTabWidget(self):
        self.tabWidget = QTabWidget()
        # self.tabWidget.setSizePolicy(QSizePolicy.Preferred,
        #         QSizePolicy.Ignored)

        self.create_phonology_tab()
        self.create_lexicon_tab()
        self.create_sound_change_tab()
        self.create_terminal_tab()
        self.create_history_tab()
        self.create_review_tab()

        self.tabWidget.addTab(self.phonologyTab, "Phonology")
        self.tabWidget.addTab(self.lexiconTab, "Lexicon")
        self.tabWidget.addTab(self.soundChangeTab, "Sound Changes")
        self.tabWidget.addTab(self.terminalTab, "Terminal")
        self.tabWidget.addTab(self.historyTab, "History")
        self.tabWidget.addTab(self.reviewTab, "Review")

    def create_phonology_tab(self):
        self.phonologyTab = QWidget()
        self.create_phoneme_class_list()
        self.create_phoneme_inventory_list()
        self.create_create_phoneme_widget()
        create_phoneme_button = QPushButton("Create phoneme")
        create_phoneme_button.pressed.connect(self.create_phoneme)

        phonologyTabVBox = QVBoxLayout()
        phonologyTabVBox.setContentsMargins(5, 5, 5, 5)

        phonologyTabHBox0 = QHBoxLayout()
        phonologyTabHBox0.addWidget(self.phoneme_class_list)
        phonologyTabHBox0.addWidget(self.phoneme_inventory_list)
        phonologyTabVBox.addLayout(phonologyTabHBox0)

        phonologyTabHBox1 = QHBoxLayout()
        phonologyTabHBox1.addWidget(create_phoneme_button)
        phonologyTabVBox.addLayout(phonologyTabHBox1)

        phonologyTabHBox2 = QHBoxLayout()
        phonologyTabHBox2.addWidget(self.create_phoneme_widget)
        phonologyTabVBox.addLayout(phonologyTabHBox2)

        self.phonologyTab.setLayout(phonologyTabVBox)

    def create_lexicon_tab(self):
        self.lexiconTab = QWidget()
        self.createLexemeList()
        self.createLexemeFormList()
        # self.clearSelectedLexeme()
        self.create_create_lexeme_widget()
        create_lexeme_button = QPushButton("Create lexeme")
        create_lexeme_button.pressed.connect(self.create_lexeme)
        create_sound_change_button = QPushButton("Create sound change")
        create_sound_change_button.pressed.connect(self.create_sound_change_from_word)
        show_word_history_button = QPushButton("Show word history")
        show_word_history_button.pressed.connect(self.show_word_history)
        # save_file_button = QPushButton("Save File")
        # save_file_button.pressed.connect(self.save_file)
        export_to_docx_button = QPushButton("Export to .DOCX")
        export_to_docx_button.pressed.connect(self.export_lexicon_to_docx)

        lexiconTabVBox = QVBoxLayout()
        lexiconTabVBox.setContentsMargins(5, 5, 5, 5)

        lexiconTabHBox0 = QHBoxLayout()
        lexiconTabHBox0.addWidget(self.lexeme_list)
        lexiconTabHBox0.addWidget(self.lexeme_form_list)
        lexiconTabVBox.addLayout(lexiconTabHBox0)

        lexiconTabHBox1 = QHBoxLayout()
        lexiconTabHBox1.addWidget(self.create_lexeme_widget)
        lexiconTabVBox.addLayout(lexiconTabHBox1)

        lexiconTabHBox2 = QHBoxLayout()
        lexiconTabHBox2.addWidget(create_lexeme_button)
        lexiconTabHBox2.addWidget(create_sound_change_button)
        lexiconTabHBox2.addWidget(show_word_history_button)
        lexiconTabVBox.addLayout(lexiconTabHBox2)

        lexiconTabHBox3 = QHBoxLayout()
        # lexiconTabHBox3.addWidget(save_file_button)
        lexiconTabHBox3.addWidget(export_to_docx_button)
        lexiconTabVBox.addLayout(lexiconTabHBox3)

        self.lexiconTab.setLayout(lexiconTabVBox)

    def create_sound_change_tab(self):
        self.soundChangeTab = QWidget()
        self.soundChangeWidget = QLineEdit()
        soundChangeLabel = QLabel("Rule")
        soundChangeLabel.setBuddy(self.soundChangeWidget)
        soundChangeGenerateButton = QPushButton("Generate rule")
        soundChangeGenerateButton.pressed.connect(self.generate_sound_change)
        applySoundChangeButton = QPushButton("Apply rule")
        applySoundChangeButton.pressed.connect(self.send_sound_change_command)
        self.fastForwardStepsInput = QLineEdit()
        fastForwardStepsLabel = QLabel("Steps to fast forward")
        fastForwardStepsLabel.setBuddy(self.fastForwardStepsInput)
        self.fastForwardButton = QPushButton("Fast forward")
        self.onlyInt = QIntValidator()
        self.fastForwardStepsInput.setValidator(self.onlyInt)
        self.fastForwardButton.pressed.connect(self.fast_forward)

        soundChangeTabVBox = QVBoxLayout()
        soundChangeTabVBox.setContentsMargins(5, 5, 5, 5)

        soundChangeTabHBox0 = QHBoxLayout()
        soundChangeTabHBox0.addWidget(soundChangeLabel)
        soundChangeTabHBox0.addWidget(self.soundChangeWidget)
        soundChangeTabVBox.addLayout(soundChangeTabHBox0)

        soundChangeTabHBox1 = QHBoxLayout()
        soundChangeTabHBox1.addWidget(soundChangeGenerateButton)
        soundChangeTabHBox1.addWidget(applySoundChangeButton)
        soundChangeTabVBox.addLayout(soundChangeTabHBox1)

        soundChangeTabHBox2 = QHBoxLayout()
        soundChangeTabHBox2.addWidget(fastForwardStepsLabel)
        soundChangeTabHBox2.addWidget(self.fastForwardStepsInput)
        soundChangeTabHBox2.addWidget(self.fastForwardButton)
        soundChangeTabVBox.addLayout(soundChangeTabHBox2)

        self.soundChangeTab.setLayout(soundChangeTabVBox)
    
    def create_terminal_tab(self):
        self.terminalTab = QWidget()
        self.terminalInputWidget = QLineEdit()
        self.terminalInputWidget.returnPressed.connect(self.send_command_to_processor)
        self.terminalOutputWidget = QTextEdit()
        self.terminalOutputWidget.setReadOnly(True)
        terminalTabVBox = QVBoxLayout()
        terminalTabVBox.setContentsMargins(5, 5, 5, 5)
        terminalTabVBox.addWidget(self.terminalOutputWidget)  # put input below output
        terminalTabVBox.addWidget(self.terminalInputWidget)
        self.terminalTab.setLayout(terminalTabVBox)

    def create_history_tab(self):
        self.historyTab = QWidget()
        historyTabVBox = QVBoxLayout()
        historyTabVBox.setContentsMargins(5, 5, 5, 5)
        historyTabVBox.addWidget(QLabel("TODO"))
        self.historyTab.setLayout(historyTabVBox)

    def create_review_tab(self):
        self.reviewTab = QWidget()
        reviewTabVBox = QVBoxLayout()
        reviewTabVBox.setContentsMargins(5, 5, 5, 5)

        sanity_check_button = QPushButton("Sanity check")
        sanity_check_button.pressed.connect(self.sanity_check)
        homophone_check_button = QPushButton("Homophone check")
        homophone_check_button.pressed.connect(self.homophone_check)
        reviewTabVBox.addWidget(sanity_check_button)
        reviewTabVBox.addWidget(homophone_check_button)
        self.reviewTab.setLayout(reviewTabVBox)

    def create_phoneme_class_list(self):
        phoneme_class_list_label = QLabel("Phoneme Classes")
        self.phoneme_class_list = QListWidget()
        phoneme_class_list_label.setBuddy(self.phoneme_class_list)
        self.phoneme_class_list.currentItemChanged.connect(self.show_phonemes_in_class)

    def create_phoneme_inventory_list(self):
        phoneme_inventory_list_label = QLabel("Phoneme Inventory")
        self.phoneme_inventory_list = QListWidget()
        phoneme_inventory_list_label.setBuddy(self.phoneme_inventory_list)

    def create_create_phoneme_widget(self):
        cpw = QWidget()
        phoneme_input_label = QLabel("Phoneme symbol (no brackets):")
        cpw.phoneme_input = QLineEdit()
        phoneme_input_label.setBuddy(cpw.phoneme_input)
        classes_input_label = QLabel("Classes this phoneme is in:")
        cpw.classes_input = QLineEdit()
        classes_input_label.setBuddy(cpw.classes_input)
        enter_button = QPushButton("Submit")
        cancel_button = QPushButton("Cancel")
        enter_button.pressed.connect(self.submit_created_phoneme)
        cancel_button.pressed.connect(lambda: self.create_phoneme_widget.hide())
        vbox = QVBoxLayout()
        vbox.addWidget(phoneme_input_label)
        vbox.addWidget(cpw.phoneme_input)
        vbox.addWidget(classes_input_label)
        vbox.addWidget(cpw.classes_input)
        vbox.addWidget(enter_button)
        vbox.addWidget(cancel_button)
        cpw.setLayout(vbox)

        self.create_phoneme_widget = cpw
        self.create_phoneme_widget.hide()


    def create_create_lexeme_widget(self):
        clw = QWidget()
        # params to get are:
        # lexeme, pos, short_gloss (optional), gloss
        # e.g. tamas = nq sand
        lexeme_input_label = QLabel("Lexeme:")
        clw.lexeme_input = QLineEdit()
        lexeme_input_label.setBuddy(clw.lexeme_input)
        pos_input_label = QLabel("POS:")
        clw.pos_input = QComboBox()  # dropdown list
        for pos in self.command_processor.get_parts_of_speech():
            clw.pos_input.addItem(pos)
        pos_input_label.setBuddy(clw.pos_input)
        short_gloss_input_label = QLabel("Short gloss (optional):")
        clw.short_gloss_input = QLineEdit()
        short_gloss_input_label.setBuddy(clw.short_gloss_input)
        gloss_input_label = QLabel("Gloss:")
        clw.gloss_input = QLineEdit()
        gloss_input_label.setBuddy(clw.gloss_input)
        enter_button = QPushButton("Submit")
        cancel_button = QPushButton("Cancel")
        enter_button.pressed.connect(self.submit_created_lexeme)
        cancel_button.pressed.connect(lambda: self.create_lexeme_widget.hide())
        vbox = QVBoxLayout()
        vbox.addWidget(lexeme_input_label)
        vbox.addWidget(clw.lexeme_input)
        vbox.addWidget(pos_input_label)
        vbox.addWidget(clw.pos_input)
        vbox.addWidget(short_gloss_input_label)
        vbox.addWidget(clw.short_gloss_input)
        vbox.addWidget(gloss_input_label)
        vbox.addWidget(clw.gloss_input)
        vbox.addWidget(enter_button)
        vbox.addWidget(cancel_button)
        clw.setLayout(vbox)

        self.create_lexeme_widget = clw
        self.create_lexeme_widget.hide()

    def submit_created_lexeme(self):
        s = ""
        s += self.create_lexeme_widget.lexeme_input.text() + " = "
        s += self.create_lexeme_widget.pos_input.currentText() + " "
        sg = self.create_lexeme_widget.short_gloss_input.text()
        if sg != "":
            s += "g:" + sg + " "
        s += self.create_lexeme_widget.gloss_input.text()
        self.command_processor.process_command(s)
        self.create_lexeme_widget.hide()

    def submit_created_phoneme(self):
        forbidden_chars = [" ", "[", "]", "\\"]
        s = "\\phone "
        phoneme_symbol = self.create_phoneme_widget.phoneme_input.text().strip()
        classes = self.create_phoneme_widget.classes_input.text().strip()
        assert all(x not in phoneme_symbol for x in forbidden_chars + list(string.ascii_uppercase))
        assert len(phoneme_symbol) > 0
        if len(phoneme_symbol) > 1:
            phoneme_symbol = "[" + phoneme_symbol + "]"
        assert " " not in classes
        classes = classes.split(",")
        new_classes = []
        for cl in classes:
            assert all(x in string.ascii_uppercase for x in cl)
            assert all(x not in cl for x in forbidden_chars)
            if len(cl) > 1:
                cl = "[" + cl + "]"
            new_classes.append(cl)
        s += "{} {}".format(phoneme_symbol, ",".join(new_classes))
        self.command_processor.process_command(s)
        self.create_phoneme_widget.hide()

        self.populate_phoneme_class_list()
        self.populate_phoneme_inventory_list()

    def add_pos(self, pos):
        self.create_lexeme_widget.pos_input.addItem(pos)
        self.create_lexeme_widget.pos_input.model().sort(0)  # sorts the dropdown list

    def populate_phoneme_class_list(self):
        self.phoneme_class_list.clear()
        phoneme_classes = self.language.get_phoneme_classes()
        for cl in phoneme_classes:
            self.phoneme_class_list.addItem(cl)

    def populate_phoneme_inventory_list(self, phoneme_class=None):
        self.phoneme_inventory_list.clear()
        phonemes = self.language.get_phonemes()
        used_phonemes = sorted(self.language.get_used_phonemes(), key=Language.unbracket_phoneme)
        unused_phonemes = sorted([x for x in phonemes if x not in used_phonemes], key=Language.unbracket_phoneme)
        if phoneme_class is not None:
            used_phonemes = [x for x in used_phonemes if x in self.language.phoneme_classes[phoneme_class]]
            unused_phonemes = [x for x in unused_phonemes if x in self.language.phoneme_classes[phoneme_class]]
        for p in used_phonemes:
            item_str = p
            self.phoneme_inventory_list.addItem(item_str)
        for p in unused_phonemes:
            item_str = p + " (unused)"
            self.phoneme_inventory_list.addItem(item_str)

    def show_phonemes_in_class(self):
        cl = self.phoneme_class_list.currentItem().text()
        self.populate_phoneme_inventory_list(phoneme_class=cl)

    def createLexemeList(self):
        self.lexeme_list = QListWidget()
        self.populateLexemeList()
        self.lexeme_list.currentItemChanged.connect(self.changeSelectedLexeme)

    def populateLexemeList(self):
        for lex in self.language.lexicon.lexemes:
            label = lex.citation_form.to_str() + " ({}, {})".format(lex.part_of_speech, lex.gloss)
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, lex)
            self.lexeme_list.addItem(item)

    def createLexemeFormList(self):
        self.lexeme_form_list = QListWidget()
        # self.lexeme_form_list.currentItemChanged.connect(self.report_current_selected_form)

    def report_current_selected_form(self):
        item = self.lexeme_form_list.currentItem()
        if item is None:
            return
        w = item.data(Qt.UserRole)
        print(w)

    def clearSelectedLexeme(self):
        self.lexeme_list.clearSelection()
        self.lexeme_form_list.clear()

    def changeSelectedLexeme(self):
        self.lexeme_form_list.clear()
        item = self.lexeme_list.currentItem()
        if item is None:
            return
        lex = item.data(Qt.UserRole)
        sorted_forms = sorted(lex.forms, key=lambda w: w.designation)
        # for form, form_gloss in lex.form_to_gloss.items():
        for form in sorted_forms:
            assert type(form) is Word
            item = QListWidgetItem(form.to_str() + " (" + form.gloss + ")")
            item.setData(Qt.UserRole, form)
            self.lexeme_form_list.addItem(item)

    def export_lexicon_to_docx(self):
        # TODO include all command lines
        # output_fp = "/home/wesley/programming/DocxTestOutput.docx"
        output_fp = "/home/wesley/programming/Language/ExamplishLexiconDocx_GENERATED.docx"
        lexicon = self.language.lexicon

        document = docx.Document()

        style = document.styles["Normal"]
        font = style.font
        font.name = "Charis SIL"
        font.size = Pt(12)

        document.add_paragraph(self.language.name + " Lexicon")
        document.add_paragraph()

        for lexeme in lexicon.lexemes:
            p = document.add_paragraph()
            p.add_run(lexeme.citation_form)
            p.add_run(" = ")
            p.add_run(lexeme.part_of_speech).italic = True
            p.add_run(" " + lexeme.gloss)

        # handle irregularities, notes, etc. after this

        document.save(output_fp)

    def generate_sound_change(self):
        rule = get_random_rules(1, self.language.lexicon.all_forms(), self.language.phoneme_classes)[0]
        self.soundChangeWidget.setText(rule.to_notation())

    def send_sound_change_command(self):
        rules = Rule.from_str(self.soundChangeWidget.text())
        for rule in rules:
            s = "\\sc " + rule.to_notation()
            self.command_processor.process_command(s)

    def apply_sound_change(self, rule):
        expanded_rules = rule.get_specific_cases(self.language.phoneme_classes, self.language.used_phonemes)
        new_lexicon = Lexicon([])
        for lexeme in self.language.lexicon.lexemes:
            new_citation_form = lexeme.citation_form.apply_rules(expanded_rules)
            new_forms = []
            for f in lexeme.forms:
                new_form = f.apply_rules(expanded_rules)
                new_forms.append(new_form)
            new_lexeme = Lexeme(new_citation_form, lexeme.part_of_speech, lexeme.gloss, forms=new_forms)
            new_lexicon.add_lexeme(new_lexeme)
        self.language.lexicon = new_lexicon
        self.update_lexicon_displays()
        self.soundChangeWidget.clear()

    def fast_forward(self):
        try:
            n_steps = int(self.fastForwardStepsInput.text())
        except ValueError:
            print("QIntValidator is not working!")
            return
        for i in range(n_steps):
            self.generate_sound_change()
            self.send_sound_change_command()

    def update_phonology_displays(self):
        self.populate_phoneme_class_list()
        self.populate_phoneme_inventory_list()

    def update_lexicon_displays(self):
        self.clearSelectedLexeme()
        self.lexeme_list.clear()
        self.lexeme_form_list.clear()
        self.populateLexemeList()

    def get_selected_word(self):
        item = self.lexeme_form_list.currentItem()
        if item is None:
            item = self.lexeme_list.currentItem()
            if item is None:
                print("no item selected")
                return None
        w = item.data(Qt.UserRole)
        assert type(w) is Word
        return w

    def create_phoneme(self):
        self.create_phoneme_widget.show()

    def create_lexeme(self):
        self.create_lexeme_widget.show()

    def create_sound_change_from_word(self):
        w = self.get_selected_word()
        if w is None:
            return
        self.tabWidget.setCurrentWidget(self.soundChangeTab)
        self.soundChangeWidget.setText(w.to_str())

    def show_word_history(self):
        w = self.get_selected_word()
        print("*??? -> {}".format(w.to_str()))
        raise NotImplementedError

    def send_command_to_processor(self):
        command_str = self.terminalInputWidget.text()
        self.command_processor.process_command(command_str)

    def update_command_history_display(self):
        self.terminalOutputWidget.setText("\n".join(self.command_processor.command_history))
        self.terminalOutputWidget.verticalScrollBar().setValue(self.terminalOutputWidget.verticalScrollBar().maximum())
        self.terminalInputWidget.clear()

    def sanity_check(self):
        forms = self.language.lexicon.all_forms()
        raise NotImplementedError  # TODO

    def homophone_check(self):
        forms = self.language.lexicon.all_forms()
        d = {}
        for w in forms:
            s = w.to_str()
            if s not in d:
                d[s] = []
            d[s].append(w)
        homophones = [s for s in d if len(d[s]) > 1]
        if len(homophones) == 0:
            print("No homophones found!")
            return
        homophones = sorted(homophones, key=lambda s: (-1*len(d[s]), s))
        print("Homophones found:")
        for s in homophones:
            meanings = [w.gloss for w in d[s]]
            print(s, "({} times)".format(len(d[s])))
            print("meanings:", meanings)


if __name__ == '__main__':
    if len(sys.argv) == 2:
        argv_input_fp = sys.argv[1]
    elif len(sys.argv) > 2:
        print("usage: python ConlangWorkspaceGUI.py (<input_filepath>)")
        sys.exit()
    else:
        argv_input_fp = None
    
    empty_lexicon = Lexicon([])
    language = Language("Examplish", empty_lexicon)

    app = QApplication(sys.argv)
    gui = ConlangWorkspaceGUI(language)
    if argv_input_fp is not None:
        gui.open_file_known_filepath(argv_input_fp)

    gui.show()
    sys.exit(app.exec_())