from PyQt5.QtCore import (QDateTime, Qt, QTimer)
from PyQt5.QtWidgets import (QAction, QApplication, QCheckBox, QComboBox, QDateTimeEdit,
        QDial, QDialog, QFileDialog, QGridLayout, QGroupBox, QHBoxLayout, 
        QInputDialog, QLabel, QLineEdit,
        QListWidget, QListWidgetItem, QMainWindow, QMenuBar,
        QProgressBar, QPushButton, QRadioButton, QScrollBar, QSizePolicy,
        QSlider, QSpinBox, QStyleFactory, QTableWidget, QTabWidget, QTextEdit,
        QVBoxLayout, QWidget
    )
from PyQt5.QtGui import (QIntValidator)

import docx
from docx.shared import Pt

from LanguageEvolutionTools import (
        Word, Rule, Lexeme, Lexicon, Language, InflectionForm, Morpheme, 
        DEFAULT_PHONEME_CLASSES, 
        evolve_word, get_random_rules, 
    )

import sys


class ConlangWorkspaceGUI(QMainWindow):
    def __init__(self, language, parent=None):
        super(ConlangWorkspaceGUI, self).__init__(parent)
        self.originalPalette = QApplication.palette()

        self.language = language
        self.command_processor = CommandProcessor(gui=self)

        self.setup_menu()

        mainLayout = QGridLayout()
        self.createTabWidget()
        mainLayout.addWidget(self.tabWidget)
        self.setCentralWidget(QWidget(self))
        self.centralWidget().setLayout(mainLayout)
        self.setWindowTitle("Conlang Workspace")

        self.setup_file_saving()
        self.setup_file_opening()

    def setup_menu(self):
        self.main_menu = self.menuBar()
        self.file_menu = self.main_menu.addMenu('File')
        #editMenu = mainMenu.addMenu('Edit')
        #viewMenu = mainMenu.addMenu('View')
        #searchMenu = mainMenu.addMenu('Search')
        #toolsMenu = mainMenu.addMenu('Tools')
        #helpMenu = mainMenu.addMenu('Help')

    def setup_file_opening(self):
        openFile = QAction("&Open File", self)
        openFile.setShortcut("Ctrl+O")
        openFile.setStatusTip('Open File')
        openFile.triggered.connect(self.open_file)
        self.file_menu.addAction(openFile)

    def open_file(self):
        fp = QFileDialog.getOpenFileName(self, 'Open File')
        if type(fp) is tuple:
            fp = fp[0]
        with open(fp) as f:
            lines = f.readlines()
        lines = [x.strip() for x in lines]
        for command in lines:
            self.command_processor.process_command(command)

    def setup_file_saving(self):
        saveFile = QAction("&Save File", self)
        saveFile.setShortcut("Ctrl+S")
        saveFile.setStatusTip('Save File')
        saveFile.triggered.connect(self.save_file)
        self.file_menu.addAction(saveFile)

    def save_file(self):
        fp = QFileDialog.getSaveFileName(self, 'Save File')
        if type(fp) is tuple:
            fp = fp[0]
        contents = ""
        for command in self.command_processor.command_history:
            contents += command + "\n"
        with open(fp, 'w') as f:
            f.write(contents)

    def createTabWidget(self):
        self.tabWidget = QTabWidget()
        # self.tabWidget.setSizePolicy(QSizePolicy.Preferred,
        #         QSizePolicy.Ignored)

        self.create_lexicon_tab()
        self.create_sound_change_tab()
        self.create_terminal_tab()
        self.create_history_tab()

        self.tabWidget.addTab(self.lexiconTab, "Lexicon")
        self.tabWidget.addTab(self.soundChangeTab, "Sound Changes")
        self.tabWidget.addTab(self.terminalTab, "Terminal")
        self.tabWidget.addTab(self.historyTab, "History")

    def create_lexicon_tab(self):
        self.lexiconTab = QWidget()
        self.createLexemeList()
        self.createLexemeFormList()
        # self.clearSelectedLexeme()
        self.create_create_lexeme_widget()
        create_lexeme_button = QPushButton("Create lexeme")
        create_lexeme_button.pressed.connect(self.create_lexeme)
        create_sound_change_button = QPushButton("Create sound change")
        create_sound_change_button.pressed.connect(self.create_sound_change_from_word)
        show_word_history_button = QPushButton("Show word history")
        show_word_history_button.pressed.connect(self.show_word_history)
        # save_file_button = QPushButton("Save File")
        # save_file_button.pressed.connect(self.save_file)
        export_to_docx_button = QPushButton("Export to .DOCX")
        export_to_docx_button.pressed.connect(self.export_lexicon_to_docx)

        lexiconTabVBox = QVBoxLayout()
        lexiconTabVBox.setContentsMargins(5, 5, 5, 5)

        lexiconTabHBox0 = QHBoxLayout()
        lexiconTabHBox0.addWidget(self.lexeme_list)
        lexiconTabHBox0.addWidget(self.lexeme_form_list)
        lexiconTabVBox.addLayout(lexiconTabHBox0)

        lexiconTabHBox1 = QHBoxLayout()
        lexiconTabHBox1.addWidget(self.create_lexeme_widget)
        lexiconTabVBox.addLayout(lexiconTabHBox1)

        lexiconTabHBox2 = QHBoxLayout()
        lexiconTabHBox2.addWidget(create_lexeme_button)
        lexiconTabHBox2.addWidget(create_sound_change_button)
        lexiconTabHBox2.addWidget(show_word_history_button)
        lexiconTabVBox.addLayout(lexiconTabHBox2)

        lexiconTabHBox3 = QHBoxLayout()
        # lexiconTabHBox3.addWidget(save_file_button)
        lexiconTabHBox3.addWidget(export_to_docx_button)
        lexiconTabVBox.addLayout(lexiconTabHBox3)
        self.lexiconTab.setLayout(lexiconTabVBox)

    def create_sound_change_tab(self):
        self.soundChangeTab = QWidget()
        self.soundChangeWidget = QLineEdit()
        soundChangeLabel = QLabel("Rule")
        soundChangeLabel.setBuddy(self.soundChangeWidget)
        soundChangeGenerateButton = QPushButton("Generate rule")
        soundChangeGenerateButton.pressed.connect(self.generate_sound_change)
        applySoundChangeButton = QPushButton("Apply rule")
        applySoundChangeButton.pressed.connect(self.apply_sound_change)
        self.fastForwardStepsInput = QLineEdit()
        fastForwardStepsLabel = QLabel("Steps to fast forward")
        fastForwardStepsLabel.setBuddy(self.fastForwardStepsInput)
        self.fastForwardButton = QPushButton("Fast forward")
        self.onlyInt = QIntValidator()
        self.fastForwardStepsInput.setValidator(self.onlyInt)
        self.fastForwardButton.pressed.connect(self.fast_forward)

        soundChangeTabVBox = QVBoxLayout()
        soundChangeTabVBox.setContentsMargins(5, 5, 5, 5)

        soundChangeTabHBox0 = QHBoxLayout()
        soundChangeTabHBox0.addWidget(soundChangeLabel)
        soundChangeTabHBox0.addWidget(self.soundChangeWidget)
        soundChangeTabVBox.addLayout(soundChangeTabHBox0)

        soundChangeTabHBox1 = QHBoxLayout()
        soundChangeTabHBox1.addWidget(soundChangeGenerateButton)
        soundChangeTabHBox1.addWidget(applySoundChangeButton)
        soundChangeTabVBox.addLayout(soundChangeTabHBox1)

        soundChangeTabHBox2 = QHBoxLayout()
        soundChangeTabHBox2.addWidget(fastForwardStepsLabel)
        soundChangeTabHBox2.addWidget(self.fastForwardStepsInput)
        soundChangeTabHBox2.addWidget(self.fastForwardButton)
        soundChangeTabVBox.addLayout(soundChangeTabHBox2)

        self.soundChangeTab.setLayout(soundChangeTabVBox)
    
    def create_terminal_tab(self):
        self.terminalTab = QWidget()
        self.terminalInputWidget = QLineEdit()
        self.terminalInputWidget.returnPressed.connect(self.send_command_to_processor)
        self.terminalOutputWidget = QTextEdit()
        self.terminalOutputWidget.setReadOnly(True)
        terminalTabVBox = QVBoxLayout()
        terminalTabVBox.setContentsMargins(5, 5, 5, 5)
        terminalTabVBox.addWidget(self.terminalOutputWidget)  # put input below output
        terminalTabVBox.addWidget(self.terminalInputWidget)
        self.terminalTab.setLayout(terminalTabVBox)

    def create_history_tab(self):
        self.historyTab = QWidget()
        historyTabVBox = QVBoxLayout()
        historyTabVBox.setContentsMargins(5, 5, 5, 5)
        historyTabVBox.addWidget(QLabel("TODO"))
        self.historyTab.setLayout(historyTabVBox)

    def create_create_lexeme_widget(self):
        clw = QWidget()
        # params to get are:
        # lexeme, pos, short_gloss (optional), gloss
        # e.g. tamas = nq sand
        lexeme_input_label = QLabel("Lexeme:")
        clw.lexeme_input = QLineEdit()
        lexeme_input_label.setBuddy(clw.lexeme_input)
        pos_input_label = QLabel("POS:")
        clw.pos_input = QComboBox()  # dropdown list
        for pos in self.command_processor.get_parts_of_speech():
            clw.pos_input.addItem(pos)
        pos_input_label.setBuddy(clw.pos_input)
        short_gloss_input_label = QLabel("Short gloss (optional):")
        clw.short_gloss_input = QLineEdit()
        short_gloss_input_label.setBuddy(clw.short_gloss_input)
        gloss_input_label = QLabel("Gloss:")
        clw.gloss_input = QLineEdit()
        gloss_input_label.setBuddy(clw.gloss_input)
        enter_button = QPushButton("Submit")
        cancel_button = QPushButton("Cancel")
        enter_button.pressed.connect(self.submit_created_lexeme)
        cancel_button.pressed.connect(lambda: self.create_lexeme_widget.hide())
        vbox = QVBoxLayout()
        vbox.addWidget(lexeme_input_label)
        vbox.addWidget(clw.lexeme_input)
        vbox.addWidget(pos_input_label)
        vbox.addWidget(clw.pos_input)
        vbox.addWidget(short_gloss_input_label)
        vbox.addWidget(clw.short_gloss_input)
        vbox.addWidget(gloss_input_label)
        vbox.addWidget(clw.gloss_input)
        vbox.addWidget(enter_button)
        vbox.addWidget(cancel_button)
        clw.setLayout(vbox)

        self.create_lexeme_widget = clw
        self.create_lexeme_widget.hide()

    def submit_created_lexeme(self):
        s = ""
        s += self.create_lexeme_widget.lexeme_input.text() + " = "
        s += self.create_lexeme_widget.pos_input.currentText() + " "
        sg = self.create_lexeme_widget.short_gloss_input.text()
        if sg != "":
            s += "g:" + sg + " "
        s += self.create_lexeme_widget.gloss_input.text()
        self.command_processor.process_command(s)
        self.create_lexeme_widget.hide()

    def add_pos(self, pos):
        self.create_lexeme_widget.pos_input.addItem(pos)
        self.create_lexeme_widget.pos_input.model().sort(0)  # sorts the dropdown list

    def createLexemeList(self):
        self.lexeme_list = QListWidget()
        self.populateLexemeList()
        self.lexeme_list.currentItemChanged.connect(self.changeSelectedLexeme)

    def populateLexemeList(self):
        for lex in self.language.lexicon.lexemes:
            label = lex.citation_form.to_str() + " ({}, {})".format(lex.part_of_speech, lex.gloss)
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, lex)
            self.lexeme_list.addItem(item)

    def createLexemeFormList(self):
        self.lexeme_form_list = QListWidget()
        # self.lexeme_form_list.currentItemChanged.connect(self.report_current_selected_form)

    def report_current_selected_form(self):
        item = self.lexeme_form_list.currentItem()
        if item is None:
            return
        w = item.data(Qt.UserRole)
        print(w)

    def clearSelectedLexeme(self):
        self.lexeme_list.clearSelection()
        self.lexeme_form_list.clear()

    def changeSelectedLexeme(self):
        self.lexeme_form_list.clear()
        item = self.lexeme_list.currentItem()
        if item is None:
            return
        lex = item.data(Qt.UserRole)
        sorted_forms = sorted(lex.forms, key=lambda w: w.designation)
        # for form, form_gloss in lex.form_to_gloss.items():
        for form in sorted_forms:
            assert type(form) is Word
            item = QListWidgetItem(form.to_str() + " (" + form.gloss + ")")
            item.setData(Qt.UserRole, form)
            self.lexeme_form_list.addItem(item)

    def export_lexicon_to_docx(self):
        # TODO include all command lines
        # output_fp = "/home/wesley/programming/DocxTestOutput.docx"
        output_fp = "/home/wesley/programming/Language/ExamplishLexiconDocx_GENERATED.docx"
        lexicon = self.language.lexicon

        document = docx.Document()

        style = document.styles["Normal"]
        font = style.font
        font.name = "Charis SIL"
        font.size = Pt(12)

        document.add_paragraph(self.language.name + " Lexicon")
        document.add_paragraph()

        for lexeme in lexicon.lexemes:
            p = document.add_paragraph()
            p.add_run(lexeme.citation_form)
            p.add_run(" = ")
            p.add_run(lexeme.part_of_speech).italic = True
            p.add_run(" " + lexeme.gloss)

        # handle irregularities, notes, etc. after this

        document.save(output_fp)

    def generate_sound_change(self):
        rule = get_random_rules(1, self.language.lexicon.all_forms(), self.language.phoneme_classes)[0]
        self.soundChangeWidget.setText(rule.to_notation())

    def apply_sound_change(self):
        rules = Rule.from_str(self.soundChangeWidget.text())#.replace("Ø", ""))  # shouldn't have null symbols anymore
        expanded_rules = []
        for rule in rules:
            expanded_rules += rule.get_specific_cases(self.language.phoneme_classes, self.language.used_phonemes)
        new_lexicon = Lexicon([])
        for lexeme in self.language.lexicon.lexemes:
            new_citation_form = evolve_word(lexeme.citation_form, expanded_rules)
            new_forms = []
            for f in lexeme.forms:
                new_form = evolve_word(f, expanded_rules)
                new_forms.append(new_form)
            new_lexeme = Lexeme(new_citation_form, lexeme.part_of_speech, lexeme.gloss, forms=new_forms)
            new_lexicon.add_lexeme(new_lexeme)
        self.language.lexicon = new_lexicon
        self.update_lexicon_displays()
        self.soundChangeWidget.clear()
        print("sound change(s) applied: {}".format(rules))

    def fast_forward(self):
        try:
            n_steps = int(self.fastForwardStepsInput.text())
        except ValueError:
            print("QIntValidator is not working!")
            return
        for i in range(n_steps):
            self.generate_sound_change()
            self.apply_sound_change()

    def update_lexicon_displays(self):
        self.clearSelectedLexeme()
        self.lexeme_list.clear()
        self.lexeme_form_list.clear()
        self.populateLexemeList()

    def get_selected_word(self):
        item = self.lexeme_form_list.currentItem()
        if item is None:
            item = self.lexeme_list.currentItem()
            if item is None:
                print("no item selected")
                return None
        w = item.data(Qt.UserRole)
        assert type(w) is Word
        return w

    def create_lexeme(self):
        self.create_lexeme_widget.show()

    def create_sound_change_from_word(self):
        w = self.get_selected_word()
        if w is None:
            return
        self.tabWidget.setCurrentWidget(self.soundChangeTab)
        self.soundChangeWidget.setText(w.to_str())

    def show_word_history(self):
        w = self.get_selected_word()
        print("*??? -> {}".format(w.to_str()))
        raise NotImplementedError

    def send_command_to_processor(self):
        command_str = self.terminalInputWidget.text()
        self.command_processor.process_command(command_str)

    def update_command_history_display(self):
        self.terminalOutputWidget.setText("\n".join(self.command_processor.command_history))
        self.terminalOutputWidget.verticalScrollBar().setValue(self.terminalOutputWidget.verticalScrollBar().maximum())
        self.terminalInputWidget.clear()


class CommandProcessor:
    def __init__(self, gui):
        self.command_history = []
        self.gui = gui

        # language-related variables for the processor to keep track of
        # maybe should move these to be attributes of the language, as long as processor can see them somehow
        self.full_inflections_by_part_of_speech = {}
        self.affixes_by_part_of_speech_and_feature = {}
        self.inflection_hierarchy_by_part_of_speech = {}
        self.inflection_templates_by_part_of_speech = {}

    @staticmethod
    def is_command_entry(s):
        return s[0] == "\\"

    @staticmethod
    def is_lexeme_entry(s):
        return s[0] not in ["\\", "#"] and " = " in s
        # use backslash to start special lines, such as defining a word class
        # use '#' for comment lines
    
    def process_command(self, command_str):
        command_str = command_str.strip()
        # if command_str == "" or command_str[0] == "#":
        #     # do nothing on blank lines and comments
        #     return
        # print("command to process:", command_str)

        if command_str == "" or command_str[0] == "#":
            pass
        elif CommandProcessor.is_command_entry(command_str):
            self.process_command_entry(command_str)
        elif CommandProcessor.is_lexeme_entry(command_str):
            self.process_lexeme_entry(command_str)
            # self.expand_templates_for_new_lexeme()
        else:
            print("command cannot be processed: {}".format(command_str))
            return

        # update the lexicon based on any new features, new affixes for a feature, new lexemes, etc.
        # self.expand_templates()  # crashes computer!

        # if command was processed successfully
        self.command_history.append(command_str)
        self.gui.update_command_history_display()

    def process_command_entry(self, ce):
        command, *rest = ce.split(" ")
        assert command[0] == "\\"
        command = command[1:]
        # print("got command {} with args {}".format(command, rest))
        if command == "pos":
            self.process_pos_command_entry(rest)
            # don't expand templates here because only declaring new pos, no inflections yet

        elif command == "inflect":
            self.process_inflect_command_entry(rest)
            # templates for this pos will be expanded in the process_inflect_... method
        
        else:
            print("unknown command \'{}\'".format(command))

    def process_pos_command_entry(self, args):
        # e.g. \pos v {negation}_{tense}{person}{number}
        #      \pos pa _
        #      \pos pb {motion}_{number}
        pos, template = args
        assert pos not in self.full_inflections_by_part_of_speech, "already have inflection template for pos \"{}\"".format(pos)
        self.gui.add_pos(pos)
        self.full_inflections_by_part_of_speech[pos] = []
        self.affixes_by_part_of_speech_and_feature[pos] = {}
        self.inflection_hierarchy_by_part_of_speech[pos] = []
        self.inflection_templates_by_part_of_speech[pos] = template
        if "{" in template:
            s1s = template.split("{")[1:]
            s2s = [s.split("}")[0] for s in s1s]
            # print("got features {}".format(s2s))
            self.affixes_by_part_of_speech_and_feature[pos] = {f: [] for f in s2s}
            # print(self.affixes_by_part_of_speech_and_feature)

    def process_inflect_command_entry(self, args):
        # e.g. \inflect v ra = negation NEG
        #      \inflect v ku = person 1
        #      \inflect v \null = number SG
        pos, morpheme_string, equals_sign, feature, gloss = args
        if morpheme_string == "\\null":
            morpheme_string = ""
        # assert pos in self.full_inflections_by_part_of_speech and len(self.full_inflections_by_part_of_speech[pos]) > 0
        assert "\\" not in morpheme_string, "non-null morpheme cannot contain backslash, but this one does: {}".format(morpheme_string)
        assert equals_sign == "="
        morpheme = Morpheme(pos, feature, morpheme_string, gloss)
        try:
            self.affixes_by_part_of_speech_and_feature[pos][feature].append(morpheme)
            if feature not in self.inflection_hierarchy_by_part_of_speech[pos]:
                self.inflection_hierarchy_by_part_of_speech[pos].append(feature)
        except KeyError:
            print("can't access index [\"{}\"][\"{}\"] in the affix dict:\n{}"
                .format(pos, feature, self.affixes_by_part_of_speech_and_feature))
            return

        self.expand_templates_for_new_affix(morpheme)
        lexemes_of_pos = [lex for lex in self.gui.language.lexicon.lexemes if lex.part_of_speech == pos]
        inflection_forms = self.full_inflections_by_part_of_speech.get(pos, [])
        for lex in lexemes_of_pos:
            lex.create_forms(inflection_forms)
        self.gui.update_lexicon_displays()

    def process_lexeme_entry(self, le):
        # e.g. lahas = n mountain
        #      mak = v eat
        #      mo = pb g:out out, outside, out of
        try:
            citation_form, rest = le.split(" = ")
            le_i = self.gui.language.lexicon.next_lexeme_designation
            citation_form = Word.from_str(citation_form, designation=str(le_i))
            pos, *gloss = rest.split(" ")
            gloss = " ".join(gloss)
            assert pos in self.full_inflections_by_part_of_speech, "unknown part of speech \"{}\"; please declare it".format(pos)
            inflection_forms = self.full_inflections_by_part_of_speech.get(pos, [])
            lexeme = Lexeme(citation_form, pos, gloss, inflection_forms=inflection_forms)
            self.gui.language.lexicon.add_lexeme(lexeme)
            self.gui.update_lexicon_displays()
        except Exception as exc:
            print("This line does not appear to be valid: {}\nIt threw {}: {}".format(le, type(exc), exc))
            raise exc

    def get_parts_of_speech(self):
        return sorted(self.full_inflections_by_part_of_speech.keys())

    def expand_templates(self):
        raise Exception("do not use!!!")
        #for pos in self.affixes_by_part_of_speech_and_feature:


    def expand_templates_for_new_lexeme(self, lexeme):
        # is this even necessary? the lexeme will get its InflectionForms by looking up those in the full_inflections_... dict
        raise NotImplementedError

    def expand_templates_for_new_affix(self, morpheme):
        pos = morpheme.pos
        template = self.inflection_templates_by_part_of_speech[pos]
        # inflections = self.full_inflections_by_part_of_speech[pos]
        inflections = [InflectionForm(pos, template, gloss="", designation_suffix="")]
        for feature in self.inflection_hierarchy_by_part_of_speech[pos]:
            morphemes = self.affixes_by_part_of_speech_and_feature[pos][feature]
            assert type(morphemes) is list, morphemes
            new_inflections = []
            for inf in inflections:
                for morpheme_i, m in enumerate(morphemes):
                    assert type(m) is Morpheme, m
                    new_morpheme_string = inf.string.replace("{"+feature+"}", m.string)
                    new_gloss = inf.gloss + "-" + m.gloss
                    new_designation_suffix = inf.designation_suffix + "." + str(morpheme_i)
                    new_inflections.append(InflectionForm(pos, new_morpheme_string, new_gloss, new_designation_suffix))
            inflections = new_inflections
        self.full_inflections_by_part_of_speech[pos] = inflections


def load_input_file(input_fp, command_processor):
    extension = input_fp.split(".")[-1]
    if extension == "docx":
        return load_lexicon_from_docx(input_fp, command_processor)
    else:
        with open(input_fp) as f:
            lines = f.readlines()
        lines = [x.strip() for x in lines]
        for s in lines:
            command_processor.process_command(s)
        return command_processor.gui.language.lexicon


def load_lexicon_from_docx(fp, command_processor):
    document = docx.Document(fp)
    ps = [p.text for p in document.paragraphs]
    for s in ps:
        command_processor.process_command(s)
    return command_processor.gui.language.lexicon



if __name__ == '__main__':
    empty_lexicon = Lexicon([])
    language = Language("Examplish", empty_lexicon, DEFAULT_PHONEME_CLASSES)

    app = QApplication(sys.argv)
    gui = ConlangWorkspaceGUI(language)

    # input_fp = "/home/wesley/programming/Language/ExamplishLexiconDocx.docx"
    input_fp = "/home/wesley/programming/Language/Examplish.cwg"
    if input("load commands/lexicon from this file?\n{}\ny/n (default yes) ".format(input_fp)) != "n":
        lexicon = load_input_file(input_fp, gui.command_processor)
        gui.language.lexicon = lexicon

    gui.show()
    sys.exit(app.exec_())