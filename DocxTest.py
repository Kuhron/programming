from docx import Document
from docx.shared import Pt


document = Document("/home/wesley/Desktop/Construction/Conlanging/Daellic/Daool Lexicon Working Version Python Test.docx")

style = document.styles["Normal"]
font = style.font
font.name = "Charis SIL"
font.size = Pt(12)

p = document.add_paragraph("Test ")
p.add_run("bold").bold = True
p.add_run(" and ")
p.add_run("italic").italic = True

p = document.add_paragraph("New paragraph")

document.save("/home/wesley/programming/DocxTestOutput.docx")